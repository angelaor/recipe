from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Helper: set paragraph shading ──
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Helper: set cell shading ──
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: set cell borders ──
def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    val.get('sz', '4'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

# ── Helper: heading ──
def add_heading(text, level=1, color_hex=None):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        if color_hex:
            run.font.color.rgb = RGBColor.from_string(color_hex)
        run.font.name = 'Calibri'
    return para

# ── Helper: body paragraph ──
def add_body(text, bold=False, color_hex=None, size=11):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return para

# ── Helper: colored section banner ──
def add_banner(text, bg_hex, fg_hex='FFFFFF'):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(f'  {text}  ')
    run.font.bold  = True
    run.font.size  = Pt(13)
    run.font.name  = 'Calibri'
    run.font.color.rgb = RGBColor.from_string(fg_hex)
    shade_paragraph(para, bg_hex)
    return para

# ── Helper: simple table ──
def add_table(headers, rows, col_widths, header_bg='0072B2', header_fg='FFFFFF'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = Inches(w)
        shade_cell(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.name  = 'Calibri'
        run.font.color.rgb = RGBColor.from_string(header_fg)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = 'EBF3FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            cell.width = Inches(w)
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                # (text, bold, color)
                run = p.add_run(val[0])
                run.font.bold  = val[1] if len(val) > 1 else False
                run.font.size  = Pt(10)
                run.font.name  = 'Calibri'
                if len(val) > 2:
                    run.font.color.rgb = RGBColor.from_string(val[2])
            else:
                run = p.add_run(str(val))
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
    return table

# ── Helper: code/monospace block ──
def add_code(lines, bg='F0F4F8'):
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.left_indent  = Inches(0.3)
        run = para.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        shade_paragraph(para, bg)

# ════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run('Outlook Email Organization')
r.font.size  = Pt(28)
r.font.bold  = True
r.font.name  = 'Calibri'
r.font.color.rgb = RGBColor.from_string('0072B2')

cover2 = doc.add_paragraph()
cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cover2.add_run('Complete Setup Reference Guide')
r2.font.size  = Pt(18)
r2.font.name  = 'Calibri'
r2.font.color.rgb = RGBColor.from_string('004E7C')

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('Folder Structure  •  Inbox Rules  •  Color Category System')
rs.font.size  = Pt(12)
rs.font.name  = 'Calibri'
rs.font.color.rgb = RGBColor.from_string('5a6270')

doc.add_paragraph()
shade_paragraph(doc.add_paragraph(''), 'E6EEF7')

doc.add_page_break()

# ════════════════════════════════════════════════
#  SECTION 1 — FOLDER STRUCTURE
# ════════════════════════════════════════════════
add_banner('SECTION 1 — Folder Structure', '0072B2')
doc.add_paragraph()

add_heading('Recommended Folder Structure', level=2, color_hex='004E7C')
add_body('Build your folders in this hierarchy inside Outlook. Right-click Inbox > New Folder to get started.', size=11)
doc.add_paragraph()

add_code([
    'Inbox  (keep only unread / action-needed here)',
    '|',
    '+-- WORK',
    '|   +-- Projects',
    '|   |   +-- Active Projects',
    '|   |   +-- Completed Projects',
    '|   +-- Meetings & Schedules',
    '|   +-- Team & Colleagues',
    '|   +-- Reports & Documents',
    '|   +-- HR & Benefits',
    '|   +-- Manager & Leadership',
    '|',
    '+-- APPLIED JOBS',
    '|   +-- Applications Sent',
    '|   |   +-- Pending Response',
    '|   |   +-- Interview Scheduled',
    '|   |   +-- Offer Received',
    '|   |   +-- Rejected / Closed',
    '|   +-- Recruiters & Agencies',
    '|   +-- Job Alerts & Listings',
    '|   +-- References & Recommendations',
    '|   +-- Research & Company Notes',
    '|',
    '+-- FINANCE',
    '|   +-- Bills & Invoices',
    '|   +-- Bank & Statements',
    '|   +-- Receipts',
    '|',
    '+-- PERSONAL',
    '|   +-- Family',
    '|   +-- Friends',
    '|   +-- Health & Medical',
    '|',
    '+-- SUBSCRIPTIONS & NEWSLETTERS',
    '|   +-- Industry News',
    '|   +-- Shopping & Promos',
    '|',
    '+-- TO-DO & FOLLOW UP',
    '|',
    '+-- ARCHIVE',
    '    +-- Work Archive',
    '    +-- Jobs Archive',
])

doc.add_paragraph()
add_heading('How to Create Folders in Outlook Desktop', level=3, color_hex='0072B2')
add_table(
    ['Step', 'Action'],
    [
        ['1', 'Right-click Inbox in the left sidebar'],
        ['2', 'Select "New Folder"'],
        ['3', 'Type the folder name (e.g., WORK) and press Enter'],
        ['4', 'Right-click the new folder to add subfolders'],
        ['5', 'Repeat for every folder in the structure above'],
    ],
    [0.6, 5.9],
)

doc.add_page_break()

# ════════════════════════════════════════════════
#  SECTION 2 — OUTLOOK RULES
# ════════════════════════════════════════════════
add_banner('SECTION 2 — Outlook Inbox Rules', '004E7C')
doc.add_paragraph()

add_heading('How to Create a Rule', level=2, color_hex='004E7C')
add_table(
    ['Step', 'Outlook Desktop', 'Outlook Web'],
    [
        ['Open Rules', 'Home tab > Rules > Manage Rules & Alerts', 'Settings (gear) > View all Outlook settings > Mail > Rules'],
        ['New Rule', 'Click "New Rule" > "Apply rule on messages I receive"', 'Click "+ Add new rule"'],
        ['Set Condition', 'Check condition > click underlined link > enter keywords', 'Choose "From" or "Subject includes" and enter value'],
        ['Set Action', 'Check "move it to specified folder" > click "specified"', 'Choose "Move to" and select your folder'],
        ['Save', 'Next > Next > name the rule > Finish', 'Click Save'],
    ],
    [0.8, 3.1, 2.6],
)

doc.add_paragraph()

# Rules list
rules = [
    ('Job Alerts & Listings', 'Applied Jobs > Job Alerts & Listings',
     'From contains', '@indeed.com | @linkedin.com | @ziprecruiter.com | @glassdoor.com | @careerbuilder.com | @monster.com'),
    ('Recruiters & Agencies', 'Applied Jobs > Recruiters & Agencies',
     'Subject contains', '"recruiter" | "talent acquisition" | "hiring manager" | "staffing" | "job opportunity" | "I came across your profile"'),
    ('Interview Scheduled', 'Applied Jobs > Applications Sent > Interview Scheduled',
     'Subject contains', '"interview" | "interview confirmation" | "schedule a call" | "calendar invite" | "let\'s connect"'),
    ('Offer Received', 'Applied Jobs > Applications Sent > Offer Received',
     'Subject contains', '"offer letter" | "job offer" | "offer of employment" | "congratulations" | "welcome to the team"'),
    ('Rejected / Closed', 'Applied Jobs > Applications Sent > Rejected / Closed',
     'Subject contains', '"not moving forward" | "other candidates" | "position has been filled" | "we regret" | "unfortunately"'),
    ('Team & Colleagues', 'Work > Team & Colleagues',
     'From contains', '@yourcompany.com  (replace with your actual work domain)'),
    ('Meetings & Schedules', 'Work > Meetings & Schedules',
     'Subject contains', '"meeting" | "calendar" | "agenda" | "invite" | "standup" | "sync" | "zoom" | "teams meeting"'),
    ('HR & Benefits', 'Work > HR & Benefits',
     'Subject contains', '"HR" | "human resources" | "payroll" | "benefits" | "PTO" | "timesheet" | "onboarding"'),
    ('Reports & Documents', 'Work > Reports & Documents',
     'Subject contains', '"report" | "document" | "please review" | "for your review" | "attached"'),
    ('Bills & Invoices', 'Finance > Bills & Invoices',
     'Subject contains', '"invoice" | "bill" | "payment due" | "statement" | "your receipt" | "amount due"'),
    ('Bank & Statements', 'Finance > Bank & Statements',
     'From contains', '@chase.com | @bankofamerica.com | @wellsfargo.com | @paypal.com  (add your bank)'),
    ('Newsletters & Promos', 'Subscriptions & Newsletters',
     'Body contains', '"unsubscribe" | "opt out" | "email preferences" | "view in browser" | "manage subscription"'),
]

for i, (name, folder, cond_type, keywords) in enumerate(rules):
    add_heading(f'Rule {i+1} — {name}', level=3, color_hex='0072B2')
    add_table(
        ['Field', 'Value'],
        [
            ['Condition', cond_type],
            ['Keywords', keywords],
            ['Action', f'Move to: {folder}'],
        ],
        [1.3, 5.2],
        header_bg='56B4E9',
        header_fg='FFFFFF',
    )
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════
#  SECTION 3 — RUN RULES ON EXISTING MAIL
# ════════════════════════════════════════════════
add_banner('SECTION 3 — Run Rules on Existing Mail', '005a50')
doc.add_paragraph()

add_heading('Outlook Desktop — Run Rules Now', level=2, color_hex='005a50')
add_table(
    ['Step', 'Action'],
    [
        ['1', 'Home tab > Rules > Manage Rules & Alerts'],
        ['2', 'Click "Run Rules Now..." at the top of the dialog'],
        ['3', 'Check all rules you want to apply (or click Select All)'],
        ['4', 'Confirm "Run in Folder" shows: Inbox'],
        ['5', 'Check "Include subfolders"'],
        ['6', 'Click "Run Now" and wait for processing to complete'],
        ['7', 'Check each folder to confirm emails sorted correctly'],
    ],
    [0.6, 5.9],
    header_bg='009E73',
)

doc.add_paragraph()
add_heading('Outlook Web — Bulk Move Existing Mail', level=2, color_hex='005a50')
add_body('Outlook Web rules only apply to NEW emails. Use these steps to clean up existing mail:', size=11)
doc.add_paragraph()

add_table(
    ['Method', 'Steps'],
    [
        ['Sort by Sender', 'Click "From" column > select group of emails > right-click > Move > choose folder'],
        ['Use Search', 'Search: from:linkedin.com > Select All (Ctrl+A) > right-click > Move > Job Alerts & Listings'],
        ['Repeat', 'Search for each sender/keyword and move to the correct folder'],
    ],
    [1.5, 5.0],
    header_bg='009E73',
)

doc.add_paragraph()
add_heading('Recommended Cleanup Order', level=3, color_hex='005a50')
add_table(
    ['Priority', 'Rule', 'Reason'],
    [
        ['1st', 'Newsletters & Promos', 'Clears the most clutter fast'],
        ['2nd', 'Job Alerts & Listings', 'High volume, easy to sort'],
        ['3rd', 'Recruiters & Agencies', 'Important to isolate'],
        ['4th', 'Interview / Offer / Rejected', 'Pipeline tracking'],
        ['5th', 'Work rules', 'Organize professional emails'],
        ['6th', 'Finance rules', 'Bills and statements'],
    ],
    [0.8, 2.5, 3.2],
    header_bg='009E73',
)

doc.add_page_break()

# ════════════════════════════════════════════════
#  SECTION 4 — COLOR CATEGORY SYSTEM
# ════════════════════════════════════════════════
add_banner('SECTION 4 — Color Category System', '7B2D8B')
doc.add_paragraph()

add_heading('Your Color-Coded Categories', level=2, color_hex='7B2D8B')
add_body('Assign these categories to emails for instant visual identification at a glance.', size=11)
doc.add_paragraph()

color_rows = [
    ('Blue',   '0072B2', 'WORK',                  'All work-related emails'),
    ('Green',  '009E73', 'APPLIED - Pending',      'Applications waiting for response'),
    ('Yellow', 'E69F00', 'APPLIED - Interview',    'Active interview emails'),
    ('Orange', 'D55E00', 'APPLIED - Offer',        'Offer letters & negotiations'),
    ('Red',    'CC3311', 'APPLIED - Rejected',     'Closed / rejected applications'),
    ('Purple', '7B2D8B', 'Recruiters',             'Recruiter & agency outreach'),
    ('Teal',   '009E73', 'Finance',                'Bills, receipts, banking'),
    ('Pink',   'CC79A7', 'Personal',               'Family & friends'),
    ('Gray',   '808080', 'Newsletters',            'Subscriptions & promos'),
    ('Gold',   'E69F00', 'FOLLOW UP',              'Action needed — any category'),
]

table = doc.add_table(rows=1 + len(color_rows), cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.LEFT
col_widths_cat = [0.8, 1.2, 1.8, 2.7]

# Header
hrow = table.rows[0]
for ci, (hdr, w) in enumerate(zip(['Color', 'Swatch', 'Category Name', 'Used For'], col_widths_cat)):
    cell = hrow.cells[ci]
    cell.width = Inches(w)
    shade_cell(cell, '7B2D8B')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor.from_string('FFFFFF')

# Data rows
for ri, (color_name, hex_c, cat_name, used_for) in enumerate(color_rows):
    row = table.rows[ri + 1]
    bg  = 'F5F0FA' if ri % 2 == 0 else 'FFFFFF'

    # Color name cell
    c0 = row.cells[0]
    c0.width = Inches(col_widths_cat[0])
    shade_cell(c0, bg)
    r0 = c0.paragraphs[0].add_run(color_name)
    r0.font.size = Pt(10); r0.font.name = 'Calibri'

    # Swatch cell — filled with the actual color
    c1 = row.cells[1]
    c1.width = Inches(col_widths_cat[1])
    shade_cell(c1, hex_c)
    r1 = c1.paragraphs[0].add_run('')
    r1.font.size = Pt(10)

    # Category name
    c2 = row.cells[2]
    c2.width = Inches(col_widths_cat[2])
    shade_cell(c2, bg)
    r2 = c2.paragraphs[0].add_run(cat_name)
    r2.font.bold = True; r2.font.size = Pt(10); r2.font.name = 'Calibri'

    # Used for
    c3 = row.cells[3]
    c3.width = Inches(col_widths_cat[3])
    shade_cell(c3, bg)
    r3 = c3.paragraphs[0].add_run(used_for)
    r3.font.size = Pt(10); r3.font.name = 'Calibri'

doc.add_paragraph()
add_heading('How to Create Categories in Outlook Desktop', level=3, color_hex='7B2D8B')
add_table(
    ['Step', 'Action'],
    [
        ['1', 'Home tab > Categorize > All Categories'],
        ['2', 'Click "New..."'],
        ['3', 'Type the category name (e.g., WORK)'],
        ['4', 'Choose the matching color from the dropdown'],
        ['5', 'Click OK and repeat for each category'],
    ],
    [0.6, 5.9],
    header_bg='7B2D8B',
)

doc.add_paragraph()
add_heading('Add Categories to Your Rules', level=3, color_hex='7B2D8B')
add_table(
    ['Rule', 'Category to Assign'],
    [
        ['Job Alerts & Listings',   'APPLIED - Pending  (Green)'],
        ['Recruiters & Agencies',   'Recruiters  (Purple)'],
        ['Interview Scheduled',     'APPLIED - Interview  (Yellow)'],
        ['Offer Received',          'APPLIED - Offer  (Orange)'],
        ['Rejected / Closed',       'APPLIED - Rejected  (Red)'],
        ['Team & Colleagues',       'WORK  (Blue)'],
        ['Meetings & Schedules',    'WORK  (Blue)'],
        ['HR & Benefits',           'WORK  (Blue)'],
        ['Bills & Invoices',        'Finance  (Teal)'],
        ['Newsletters & Promos',    'Newsletters  (Gray)'],
    ],
    [2.5, 4.0],
    header_bg='7B2D8B',
)

doc.add_paragraph()
add_heading('Keyboard Shortcuts for Quick Categorizing', level=3, color_hex='7B2D8B')
add_body('Assign shortcut keys: Home > Categorize > All Categories > click a category > set Shortcut Key', size=11)
doc.add_paragraph()
add_table(
    ['Shortcut Key', 'Category'],
    [
        ['F2', 'WORK  (Blue)'],
        ['F3', 'APPLIED - Pending  (Green)'],
        ['F4', 'APPLIED - Interview  (Yellow)'],
        ['F5', 'APPLIED - Offer  (Orange)'],
        ['F6', 'FOLLOW UP  (Gold)'],
    ],
    [1.5, 5.0],
    header_bg='7B2D8B',
)

doc.add_page_break()

# ════════════════════════════════════════════════
#  SECTION 5 — QUICK REFERENCE CHEAT SHEET
# ════════════════════════════════════════════════
add_banner('SECTION 5 — Quick Reference Cheat Sheet', 'E69F00', '1a1a1a')
doc.add_paragraph()

add_heading('Applied Jobs Pipeline at a Glance', level=2, color_hex='004E7C')
add_table(
    ['Color', 'Status', 'Folder', 'Next Action'],
    [
        ('Green',  'Pending Response',    'Applications Sent > Pending',     'Follow up after 1 week'),
        ('Yellow', 'Interview Scheduled', 'Applications Sent > Interview',   'Prepare & confirm details'),
        ('Orange', 'Offer Received',      'Applications Sent > Offer',       'Review, negotiate, decide'),
        ('Red',    'Rejected / Closed',   'Applications Sent > Rejected',    'Note feedback, move on'),
        ('Purple', 'Recruiter Contact',   'Recruiters & Agencies',           'Respond within 24 hours'),
    ],
    [0.9, 1.8, 2.3, 1.5],
    header_bg='0072B2',
)

doc.add_paragraph()
add_heading('Daily Inbox Habits', level=2, color_hex='004E7C')
add_table(
    ['Habit', 'Benefit'],
    [
        ('Process inbox once daily',           'Prevents backlog buildup'),
        ('Apply color category when reading',  'Visual priority at a glance'),
        ('Archive, never delete',              'Keeps history searchable'),
        ('Flag anything needing reply',        'Nothing falls through the cracks'),
        ('Review Applied Jobs weekly',         'Stay on top of your job pipeline'),
        ('Empty Deleted Items monthly',        'Keeps mailbox size in check'),
    ],
    [2.8, 3.7],
    header_bg='0072B2',
)

doc.add_paragraph()
add_heading('Pro Tips', level=2, color_hex='004E7C')

tips = [
    'Order matters — put specific rules above general ones in the Rules Manager to avoid conflicts.',
    'Use "Stop processing more rules" on each rule to prevent one email going to multiple folders.',
    'Search operators: from:name@email.com or subject:"interview" to find emails quickly.',
    'View > Arrangement > Categories to group your entire inbox by color at any time.',
    'Right-click any email > Sweep to bulk-delete all old emails from a sender in one click.',
]
for tip in tips:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(tip)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

# ── Footer ──
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(footer_para, 'E6EEF7')
fr = footer_para.add_run('Outlook Email Organization Guide  |  Keep this guide handy while setting up your inbox!')
fr.font.size  = Pt(10)
fr.font.name  = 'Calibri'
fr.font.color.rgb = RGBColor.from_string('004E7C')

# ── Save ──
out_path = r'C:\Users\angel\Documents\cc\recipe\Outlook_Email_Organization_Guide.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
